# Horoscope teller
import requests
import docx
from bs4 import BeautifulSoup
from datetime import date


# uses your month and day birth to evaluate your zodiac
def zodiac_by_birth_date(month: int, day: int) -> str:
    my_birth_date = date(2000, month, day)
    zodiac_signs = {
        'baran': [date(2000, 3, 21), date(2000, 4, 19)],
        'byk': [date(2000, 4, 20), date(2000, 5, 22)],
        'bliźnięta': [date(2000, 5, 23), date(2000, 6, 21)],
        'rak': [date(2000, 6, 22), date(2000, 7, 22)],
        'lew': [date(2000, 7, 23), date(2000, 8, 23)],
        'panna': [date(2000, 8, 28), date(2000, 9, 22)],
        'waga': [date(2000, 9, 23), date(2000, 10, 22)],
        'skorpion': [date(2000, 10, 23), date(2000, 11, 21)],
        'strzelec': [date(2000, 11, 22), date(2000, 12, 21)],
        'koziorożec': [date(2000, 12, 22), date(2000, 1, 19)],
        'wodnik': [date(2000, 1, 20), date(2000, 2, 18)],
        'ryby': [date(2000, 2, 19), date(2000, 3, 20)]
    }

    for key, value in zodiac_signs.items():
        if value[0] <= my_birth_date <= value[1]:
            return key


def horoscope_parsing(month: int, day: int):
    my_zodiac = zodiac_by_birth_date(month, day)
    ratings = []

# getting horoscope url based on birth date
    r = requests.get("http://horoskopy.gazeta.pl/horoskop/0,0.html")
    soup = BeautifulSoup(r.text, 'html.parser')
    yearly_horoscope_url = soup.find('li', attrs={'class': my_zodiac}).find('ul', attrs={'class': 'p1'}).find_all('li')[-1].find('a')['href']
    r = requests.get("http://horoskopy.gazeta.pl" + yearly_horoscope_url)
    soup = BeautifulSoup(r.text, 'html.parser')

# scraping all components of horoscope
    horoscope_in_html = soup.find('section', attrs={'class': 'mod mod_spec_sign'})

    # creating list of object in main horoscope text
    main_text_in_html = horoscope_in_html.find('p', attrs={'class': 'lead'}).contents

    # main horoscope information scraping
    title = horoscope_in_html.find('h1', attrs={"class": 'title'}).text
    year = horoscope_in_html.find('div', attrs={"class": 'date-container'}).contents[0]
    year = year[:4]

    ratings_lines = horoscope_in_html.find('ul', attrs={"class": 'h_rate'}).find_all('li')
    for line in ratings_lines:
        ratings.append(line.find('div').text + ' ' + line.contents[1]['class'][1][-1])

    return title, year, ratings, main_text_in_html


def creating_horoscope_file(title, year, ratings, main_text_in_html) -> None:
    # opening document and creating document variable
    document = docx.Document()
    # adding document title
    title_paragraph = document.add_paragraph(title, style='Title')
    title_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    # adding horoscope year
    document.add_paragraph(year, style='Subtitle')
    # adding horoscope ratings
    ratings_paragraph = document.add_paragraph(style='Subtitle')
    [ratings_paragraph.add_run(rating + '\n') for rating in ratings]
    # adding main text
    for text in main_text_in_html:
        if str(text) == '<br/>':
            continue
        else:
            if str(text)[0:3] == '<b>':
                document.add_paragraph(str(text)[3:-4], style='Heading 2')
            else:
                document.add_paragraph(text, style='Normal')
    # saving document
    document.save('horoscope.docx')


def main_horoscope_app() -> None:
    month = int(input("What is your birth month: "))
    day = int(input("What is your birth day: "))

    horoscope_elements = horoscope_parsing(month, day)
    creating_horoscope_file(*horoscope_elements)


main_horoscope_app()
